from sfcbidlist import bid_list
from docx import Document
# https://python-docx.readthedocs.io/en/latest/

def main():

    # Simple print for now
    sfc_bids = bid_list()
    for task in sfc_bids:
        print(task["content"])

# MS Word Document
def numbered_list(arr):
    doc = Document()
    doc.add_paragraph('first item in ordered list', style='List Number')
    doc.add_paragraph('second item in ordered list', style='List Number')
    doc.save("test.docx")

# MS Outlook
# https://www.youtube.com/watch?v=ZvmFHwAjXHI&feature=youtu.be

if __name__ == "__main__":
    main()
